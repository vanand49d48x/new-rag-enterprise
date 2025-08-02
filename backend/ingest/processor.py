# File: backend/ingest/processor.py

import os
import logging
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
import io

# Document processing libraries
try:
    from pypdf2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logging.warning("PyPDF2 not available. PDF processing will be limited.")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX processing will be limited.")

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logging.warning("openpyxl not available. Excel processing will be limited.")

try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False
    logging.warning("LangChain not available. Using simple text splitting.")

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Processor for document files including PDF, DOCX, TXT, and other text formats.
    Handles text extraction and intelligent chunking.
    """
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = ["\n\n", "\n", ". ", "! ", "? ", " ", ""]
        
        # Initialize text splitter
        if LANGCHAIN_AVAILABLE:
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separators=self.separators
            )
        else:
            self.text_splitter = None
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from various document formats
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self._extract_from_excel(file_path)
            elif file_ext in ['.txt', '.md', '.csv']:
                return self._extract_from_text(file_path)
            else:
                # Try to read as text file
                return self._extract_from_text(file_path)
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        if not PYPDF2_AVAILABLE:
            raise RuntimeError("PyPDF2 not available for PDF processing")
        
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                text_parts = []
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_parts.append(page_text)
                    except Exception as e:
                        logger.warning(f"Error extracting text from PDF page {page_num}: {e}")
                
                return "\n\n".join(text_parts)
                
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not available for DOCX processing")
        
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise
    
    def _extract_from_excel(self, file_path: str) -> str:
        """Extract text from Excel files"""
        if not OPENPYXL_AVAILABLE:
            raise RuntimeError("openpyxl not available for Excel processing")
        
        try:
            workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = []
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell in row:
                        if cell is not None:
                            row_text.append(str(cell))
                    if row_text:
                        sheet_text.append(", ".join(row_text))
                
                if sheet_text:
                    text_parts.append(f"[Sheet: {sheet_name}]\n" + "\n".join(sheet_text))
            
            workbook.close()
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error processing Excel {file_path}: {e}")
            raise
    
    def _extract_from_text(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
                
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            raise
    
    def chunk_text(self, text: str, chunk_size: Optional[int] = None) -> List[str]:
        """
        Split text into chunks using intelligent splitting
        
        Args:
            text: Text to chunk
            chunk_size: Optional custom chunk size
            
        Returns:
            List of text chunks
        """
        if not text.strip():
            return []
        
        # Use custom chunk size if provided
        if chunk_size is not None:
            current_chunk_size = chunk_size
        else:
            current_chunk_size = self.chunk_size
        
        # Use LangChain splitter if available
        if self.text_splitter and LANGCHAIN_AVAILABLE:
            try:
                # Update chunk size if needed
                if chunk_size is not None:
                    self.text_splitter.chunk_size = chunk_size
                    self.text_splitter.chunk_overlap = min(chunk_size // 10, self.chunk_overlap)
                
                chunks = self.text_splitter.split_text(text)
                logger.info(f"Split text into {len(chunks)} chunks using LangChain")
                return chunks
                
            except Exception as e:
                logger.warning(f"LangChain splitting failed, falling back to simple splitting: {e}")
        
        # Fallback to simple recursive splitting
        return self._split_text_recursively(text, current_chunk_size)
    
    def _split_text_recursively(self, text: str, chunk_size: int) -> List[str]:
        """Simple recursive text splitting without LangChain dependency"""
        if len(text) <= chunk_size:
            return [text]
        
        # Try to split on separators in order of preference
        for separator in self.separators:
            if separator in text:
                parts = text.split(separator)
                if len(parts) > 1:
                    chunks = []
                    current_chunk = ""
                    
                    for part in parts:
                        if len(current_chunk + separator + part) <= chunk_size:
                            current_chunk += separator + part if current_chunk else part
                        else:
                            if current_chunk:
                                chunks.append(current_chunk)
                            current_chunk = part
                    
                    if current_chunk:
                        chunks.append(current_chunk)
                    
                    # Recursively split chunks that are still too long
                    final_chunks = []
                    for chunk in chunks:
                        if len(chunk) > chunk_size:
                            final_chunks.extend(self._split_text_recursively(chunk, chunk_size))
                        else:
                            final_chunks.append(chunk)
                    
                    return final_chunks
        
        # If no separator works, split by character count
        chunks = []
        overlap = min(chunk_size // 10, self.chunk_overlap)
        
        for i in range(0, len(text), chunk_size - overlap):
            chunk = text[i:i + chunk_size]
            if chunk:
                chunks.append(chunk)
        
        return chunks
    
    def process_document(self, file_path: str, chunk_size: Optional[int] = None) -> List[str]:
        """
        Complete document processing pipeline
        
        Args:
            file_path: Path to the document file
            chunk_size: Optional custom chunk size
            
        Returns:
            List of text chunks
        """
        try:
            # Extract text
            text = self.extract_text(file_path)
            
            # Chunk text
            chunks = self.chunk_text(text, chunk_size)
            
            logger.info(f"Processed document {file_path}: {len(chunks)} chunks created")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            raise
    
    def get_document_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic information about a document"""
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            stat = os.stat(file_path)
            ext = Path(file_path).suffix.lower()
            
            info = {
                "file_path": file_path,
                "file_name": os.path.basename(file_path),
                "file_size": stat.st_size,
                "extension": ext,
                "is_supported": self._is_supported_format(ext)
            }
            
            return info
            
        except Exception as e:
            logger.error(f"Error getting document info for {file_path}: {e}")
            raise
    
    def _is_supported_format(self, extension: str) -> bool:
        """Check if a file format is supported"""
        supported_formats = ['.pdf', '.docx', '.doc', '.txt', '.md', '.csv', '.xlsx', '.xls']
        return extension.lower() in supported_formats
